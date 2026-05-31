# baza_word.py — Базаның барлық кестелерін Word файлға (.docx) шығарады
# Іске қосу: python database/baza_word.py (proekt/ папкасынан)

import sqlite3
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Алдымен python-docx орнату керек:")
    print("  .venv\\Scripts\\pip install python-docx")
    sys.exit(1)

# Жолдар
db_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'news_kz.db')
output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'docs', 'baza_strukturasy.docx')

# Кесте ұяшығына фон түс беру
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# Ұяшық мәтінін баптау
def style_cell(cell, text, bold=False, font_size=10, color=None, bg=None, center=False):
    cell.text = ''
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if bg:
        set_cell_bg(cell, bg)

# Деректерді оқу
conn = sqlite3.connect(db_path)
cursor = conn.cursor()
cursor.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
tables = [t[0] for t in cursor.fetchall()]

# Әр кестеде жолдар санын алу
def get_row_count(table_name):
    try:
        cursor.execute(f"SELECT COUNT(*) FROM '{table_name}'")
        return cursor.fetchone()[0]
    except:
        return '—'

# Word документін жасау
doc = Document()

# Беттің өлшемін баптау (A4)
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# Стиль — Негізгі шрифт
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ─── Бас тақырып ───
title = doc.add_heading('', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('База данных: news_kz.db')
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run(f'Жоба: Жаңалық сайты (Қазақстан) | Кестелер саны: {len(tables)}')
sub_run.font.size = Pt(10)
sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
doc.add_paragraph()

# ─── Жалпы шолу кестесі ───
doc.add_heading('Жалпы шолу', level=1)

overview = doc.add_table(rows=1, cols=3)
overview.style = 'Table Grid'
overview.alignment = WD_TABLE_ALIGNMENT.CENTER

# Баған ендері
overview.columns[0].width = Cm(6)
overview.columns[1].width = Cm(9)
overview.columns[2].width = Cm(3)

# Тақырып жол
hdr = overview.rows[0].cells
style_cell(hdr[0], 'Кесте атауы', bold=True, bg='1A56DB', color=(255,255,255), center=True)
style_cell(hdr[1], 'Мақсаты', bold=True, bg='1A56DB', color=(255,255,255), center=True)
style_cell(hdr[2], 'Жолдар', bold=True, bg='1A56DB', color=(255,255,255), center=True)

# Кесте сипаттамалары
descriptions = {
    'user':           'Пайдаланушылар (admin, тіркелген қолданушылар)',
    'category':       'Жаңалық категориялары (спорт, саясат, т.б.)',
    'news':           'Жаңалықтар (қаз/рус/ағылшын тілдерінде)',
    'source':         'RSS дереккөздері (Tengrinews, Nur.kz, т.б.)',
    'setting':        'Сайт баптаулары (аты, логотип, т.б.)',
    'bookmark':       'Пайдаланушының бетбелгілері',
    'like':           'Жаңалықтарға қойылған лайктар',
    'comment':        'Жаңалықтарға жазылған пікірлер',
    'sqlite_sequence':'SQLite ішкі автоинкремент кестесі',
}
row_colors = ['F9FAFB', 'FFFFFF']
for i, tname in enumerate(tables):
    row = overview.add_row().cells
    style_cell(row[0], tname, bold=True, bg=row_colors[i % 2])
    style_cell(row[1], descriptions.get(tname, '—'), bg=row_colors[i % 2])
    style_cell(row[2], str(get_row_count(tname)), center=True, bg=row_colors[i % 2])

doc.add_paragraph()

# ─── Әр кесте үшін толық сипаттама ───
type_colors = {
    'INTEGER': 'E1EFFE',
    'VARCHAR': 'DEF7EC',
    'TEXT':    'FDF6B2',
    'BOOLEAN': 'FCE8F3',
    'DATETIME':'FFF3CD',
}

for tname in tables:
    if tname == 'sqlite_sequence':
        continue

    doc.add_heading(f'Кесте: {tname}', level=2)

    # Кесте туралы ақпарат
    info_para = doc.add_paragraph()
    info_para.add_run(f'Сипаттама: ').bold = True
    info_para.add_run(descriptions.get(tname, '—'))

    # Баған ақпараты
    cursor.execute(f"PRAGMA table_info('{tname}')")
    columns = cursor.fetchall()

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Баған ендері
    widths = [Cm(1.2), Cm(4.5), Cm(3.5), Cm(2), Cm(4.8)]
    for j, width in enumerate(widths):
        for cell in tbl.columns[j].cells:
            cell.width = width

    # Тақырып
    hdr_cells = tbl.rows[0].cells
    headers = ['#', 'Баған атауы', 'Деректер түрі', 'NULL?', 'Сипаттама']
    for j, h in enumerate(headers):
        style_cell(hdr_cells[j], h, bold=True, bg='374151', color=(255,255,255), center=True)

    # Баған сипаттамалары
    col_descriptions = {
        'id':              'Бірегей ID (автоматты)',
        'username':        'Пайдаланушы аты',
        'password_hash':   'Шифрланған пароль',
        'email':           'Электрондық пошта',
        'phone_number':    'Телефон нөмірі',
        'google_id':       'Google аккаунт ID',
        'avatar_url':      'Профиль суретінің сілтемесі',
        'created_at':      'Жасалған уақыты',
        'is_admin':        'Администратор ма?',
        'facebook_id':     'Facebook аккаунт ID',
        'code':            'Категория коды',
        'category_id':     'Категория ID (сілтеме)',
        'title_kk':        'Тақырып — қазақша',
        'title_ru':        'Тақырып — орысша',
        'title_en':        'Тақырып — ағылшынша',
        'content_kk':      'Мазмұн — қазақша',
        'content_ru':      'Мазмұн — орысша',
        'content_en':      'Мазмұн — ағылшынша',
        'image_filename':  'Сурет файлының аты',
        'views':           'Қаралым саны',
        'source_name':     'Дереккөз атауы',
        'original_url':    'Бастапқы жаңалық сілтемесі',
        'summary_ru':      'AI қорытынды — орысша',
        'summary_kk':      'AI қорытынды — қазақша',
        'summary_en':      'AI қорытынды — ағылшынша',
        'name':            'Атауы',
        'url':             'Сайт/RSS сілтемесі',
        'source_type':     'Дереккөз түрі',
        'is_active':       'Белсенді ме?',
        'last_fetched':    'Соңғы жаңарту уақыты',
        'language':        'Тілі (ru/kk/en)',
        'key':             'Баптау кілті',
        'value':           'Баптау мәні',
        'user_id':         'Пайдаланушы ID (сілтеме)',
        'news_id':         'Жаңалық ID (сілтеме)',
        'content':         'Пікір мазмұны',
        'guest_name':      'Қонақ пайдаланушы аты',
        'seq':             'Соңғы ID мәні',
    }

    for i, col in enumerate(columns):
        col_id, col_name, col_type, notnull, default, pk = col
        row_cells = tbl.add_row().cells
        bg = row_colors[i % 2]
        
        # Деректер түріне қарай түс
        base_type = col_type.split('(')[0].upper() if col_type else ''
        type_bg = type_colors.get(base_type, bg)

        style_cell(row_cells[0], str(col_id + 1), center=True, bg=bg)
        style_cell(row_cells[1], col_name, bold=(pk == 1), bg=bg)
        style_cell(row_cells[2], col_type or '—', center=True, bg=type_bg)
        style_cell(row_cells[3], 'ЖОҚ' if notnull else 'ИӘ', center=True, bg=bg,
                   color=(0xDC, 0x26, 0x26) if notnull else (0x05, 0x96, 0x69))
        style_cell(row_cells[4], col_descriptions.get(col_name, '—'), bg=bg)

    doc.add_paragraph()

# Жол санын жабу
conn.close()

# ─── Аяқтау жазуы ───
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run('Жасалды: Python скрипті арқылы | news_kz.db')
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

# Сақтау
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Word файлы жасалды: {output_path}")
